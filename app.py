import io
import json
import os
import re
import hashlib
import uuid
from datetime import datetime

import requests
from docx import Document
from dotenv import load_dotenv
from flask import Flask, jsonify, redirect, render_template, request, send_file, session, url_for
from werkzeug.utils import secure_filename

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, ".env"), override=True)

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret-change-me")

DATA_DIR = os.path.join(BASE_DIR, "data")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads", "payment_proofs")
USERS_FILE = os.path.join(DATA_DIR, "users.json")
PAYMENTS_FILE = os.path.join(DATA_DIR, "payments.json")

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)


def _ensure_json_file(path: str, default_obj):
    if not os.path.exists(path):
        with open(path, "w", encoding="utf-8") as f:
            json.dump(default_obj, f, ensure_ascii=True, indent=2)


_ensure_json_file(USERS_FILE, {"users": []})
_ensure_json_file(PAYMENTS_FILE, {"payments": []})


def _read_json(path: str, default_obj):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default_obj


def _write_json(path: str, payload):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=True, indent=2)


def _hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def _find_user_by_username(users: list[dict], username: str):
    normalized = username.strip().lower()
    for user in users:
        if str(user.get("username", "")).strip().lower() == normalized:
            return user
    return None


def _find_user_by_email(users: list[dict], email: str):
    normalized = email.strip().lower()
    for user in users:
        if str(user.get("email", "")).strip().lower() == normalized:
            return user
    return None


def _is_duplicate_password(users: list[dict], password_hash: str) -> bool:
    for user in users:
        if user.get("passwordHash") == password_hash:
            return True
    return False


def _allowed_payment_file(filename: str) -> bool:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in {"png", "jpg", "jpeg", "webp", "pdf"}


def _safe_next_path(value: str, default: str = "") -> str:
    candidate = (value or "").strip()
    if not candidate:
        return default

    # Allow only same-site relative paths.
    if not candidate.startswith("/"):
        return default
    if candidate.startswith("//"):
        return default

    return candidate


MARKET_RESEARCH_SYSTEM_PROMPT = """You are a senior product-focused market research analyst specializing in SaaS and tech ecosystems.

Your goal is NOT just to list competitors, but to deeply analyze the market landscape of the given product idea.

Task:
Analyze the product idea and identify highly relevant, real-world competitors including SaaS products, startups, and open-source GitHub projects.

STRICT RULES:
- Only include REAL, verifiable, and existing products
- Do NOT hallucinate or invent names, links, or companies
- If you are not confident about a product, SKIP it
- Prioritize products that are functionally similar (not just loosely related)
- Focus on products solving the SAME or CLOSELY RELATED problem
- Prefer well-known tools, funded startups, or active GitHub projects
- Avoid outdated or dead products unless highly relevant

ANALYSIS DEPTH:
For each competitor:
- Understand their CORE purpose (not surface-level)
- Focus on their MAIN value proposition
- Identify their key features (not generic ones)
- Clearly explain how they overlap with the user's idea
- Be precise in similarity scoring (0 = unrelated, 10 = nearly identical)

INPUT:
Product Idea:


Current Features:
{features}

OUTPUT REQUIREMENTS:
- Find 3–5 HIGHLY RELEVANT competitors (quality > quantity)
- Provide:
  - Product name
  - Working URL (accurate)
  - Problem they solve
  - Solution approach
  - Main features (concise and meaningful)
  - Relation to this idea (clear comparison)
  - Similarity score (0–10)

Additionally:
- Provide a concise MARKET SUMMARY explaining the current landscape
- Clearly explain the DIFFERENTIATION FACTOR (what makes this idea unique or competitive)

IMPORTANT:
Focus on DEPTH, ACCURACY, and RELEVANCE over quantity.
"""

STRUCTURE_ANALYSIS_SYSTEM_PROMPT = """You are a senior product analyst and data structuring expert.

Your goal is to convert unstructured market research into clean, accurate, and structured product intelligence.

Task:
Extract and transform the provided market research into a well-defined JSON structure.

STRICT RULES:
- Do NOT hallucinate or invent any data
- ONLY use information explicitly present in the research
- If a field is missing, infer carefully OR keep it minimal (do not fabricate)
- Ensure all URLs are EXACTLY as given in the research (no modification)
- Maintain consistency across all competitor entries
- Similarity scores must remain unchanged (do not alter values)

DATA QUALITY RULES:
- Keep text concise, clear, and professional
- Avoid vague or generic descriptions
- Normalize competitor data (same structure for all)
- Extract meaningful feature lists (avoid filler points)

INPUT:
Original Idea:


Features:
{features}

Market Research:


OUTPUT REQUIREMENTS:
- Generate a clean and professional project TITLE
- Extract a clear and focused PROBLEM statement
- Extract a strong and practical SOLUTION
- Structure all competitors into consistent format
- Provide a meaningful MARKET SUMMARY
- Provide a clear DIFFERENTIATION FACTOR

OUTPUT FORMAT:
Return ONLY valid JSON (no markdown, no explanation, no extra text)
"""

SUGGESTIONS_SYSTEM_PROMPT = """You are a senior startup innovation strategist and product thinker.

Your goal is to generate HIGH-IMPACT, NON-OBVIOUS improvements that can make a product stand out in a competitive market.

Task:
Based on the provided product idea and its competitor analysis, generate 5 powerful improvements.

STRICT RULES:
- Do NOT give generic suggestions (e.g., "improve UI", "add AI", "make it faster")
- Each suggestion must be SPECIFIC, ACTIONABLE, and PRACTICAL
- Focus on real differentiation, not minor enhancements
- Base ideas on competitor weaknesses, missing features, or unexplored opportunities
- Avoid obvious or commonly used startup ideas
- Think like a founder trying to WIN the market

QUALITY REQUIREMENTS:
- Each suggestion should feel like a strong product decision
- Prefer ideas that combine multiple concepts (feature + strategy + tech)
- Highlight opportunities competitors have missed
- Suggestions should be implementable (not futuristic fantasy)
"""

SRS_SYSTEM_PROMPT = """You are a senior product requirements writer focused on MVP planning.

Your goal is to generate a practical and build-ready MVP requirements document in plain text.

Task:
Create an MVP Requirements Document using the exact section structure provided below.

STRICT RULES:
- Follow the exact section numbering and order
- Keep the document specific to the provided product context
- Avoid filler text, generic startup advice, and repeated points
- Keep language clear, concise, and implementation-oriented
- Include concrete details for flows, requirements, and architecture choices

OUTPUT REQUIREMENTS:
- Return ONLY plain text
- Do NOT use Markdown syntax (#, ##, -, *, ```)
- Do NOT wrap output in code fences
- Use this exact template and fill each section with relevant details

MVP Requirements Document

1. Product Overview
Problem:
Solution:
MVP Goal:

2. Core Features

3. User Flow
Step-by-step product flow:

4. Functional Requirements
Inputs:
System behavior:
Outputs:

5. Basic Technical Architecture
Frontend:
Backend:
Database / APIs:

6. Non-Functional Requirements
Performance:
Security:

7. Phase Plan
Phase 1 (MVP):
Phase 2 (Future):
"""


def get_env(name: str, default: str = "") -> str:
    value = os.getenv(name, default)
    return value.strip() if value else ""


def clamp_text(text: str, max_chars: int) -> str:
    text = (text or "").strip()
    if len(text) <= max_chars:
        return text
    return f"{text[:max_chars]}\n\n[Truncated to fit request size limits.]"


def compact_features(features: list[str], max_items: int = 20, max_chars_per_item: int = 220) -> list[str]:
    compacted = []
    for item in features[:max_items]:
        compacted.append(clamp_text(str(item), max_chars_per_item))
    return [item for item in compacted if item]


def compact_competitors(competitors: list[dict], max_items: int = 8, max_chars_per_field: int = 420) -> list[dict]:
    compacted = []
    for comp in competitors[:max_items]:
        if not isinstance(comp, dict):
            continue

        compacted.append(
            {
                "name": clamp_text(str(comp.get("name", "")), max_chars_per_field),
                "url": clamp_text(str(comp.get("url", "")), max_chars_per_field),
                "similarityScore": comp.get("similarityScore", 0),
                "problem": clamp_text(str(comp.get("problem", "")), max_chars_per_field),
                "solution": clamp_text(str(comp.get("solution", "")), max_chars_per_field),
                "mainFeatures": compact_features(comp.get("mainFeatures", []), max_items=8, max_chars_per_item=140),
                "relationToIdea": clamp_text(str(comp.get("relationToIdea", "")), max_chars_per_field),
            }
        )
    return compacted


def get_client_config() -> dict:
    base_url = get_env("GROQ_BASE_URL", "https://api.groq.com/openai/v1").rstrip("/")
    api_key = get_env("GROQ_API_KEY")
    model = get_env("MODEL", "groq/compound")

    if not api_key:
        raise ValueError("Missing API key. Set GROQ_API_KEY in .env")

    if not model:
        raise ValueError("Missing model. Set MODEL in .env (example: openai/gpt-oss-120b).")

    return {
        "api_key": api_key,
        "base_url": base_url,
        "model": model,
    }


def call_chat_model(model: str, system_prompt: str, user_prompt: str, temperature: float = 0.2) -> str:
    config = get_client_config()
    base_url = config["base_url"]
    api_key = config["api_key"]

    requested_model = model.strip()
    model_candidates = [requested_model]

    # Groq accepts fully-qualified IDs like "groq/compound".
    # Keep the requested model first, then try practical fallbacks.
    if requested_model == "groq/compound":
        model_candidates.extend([
            "groq/compound-mini",
            "llama-3.3-70b-versatile",
        ])

    # De-duplicate while preserving order.
    model_candidates = list(dict.fromkeys(model_candidates))

    # Prepare progressively smaller prompt variants to recover from
    # provider-side request size limits.
    prompt_variants = [user_prompt]
    for max_chars in (12000, 8000, 5000, 3000):
        if len(user_prompt) > max_chars:
            prompt_variants.append(clamp_text(user_prompt, max_chars))
    prompt_variants = list(dict.fromkeys(prompt_variants))

    url = f"{base_url}/chat/completions"

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    last_error = None
    for model_index, candidate_model in enumerate(model_candidates):
        for prompt_index, candidate_prompt in enumerate(prompt_variants):
            payload = {
                "model": candidate_model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": candidate_prompt},
                ],
                "temperature": temperature,
            }

            response = requests.post(url, headers=headers, json=payload, timeout=90)
            if response.status_code >= 400:
                last_error = response

                error_code = ""
                error_message = ""
                try:
                    error_payload = response.json().get("error", {})
                    error_code = str(error_payload.get("code", "")).strip().lower()
                    error_message = str(error_payload.get("message", "")).strip().lower()
                except Exception:
                    pass

                is_too_large = (
                    response.status_code == 413
                    or error_code == "request_too_large"
                    or "request entity too large" in error_message
                )

                # Retry same model with a smaller prompt variant if available.
                if is_too_large and prompt_index < len(prompt_variants) - 1:
                    continue

                # Try next fallback model on client-side 400 responses.
                if response.status_code == 400 and model_index < len(model_candidates) - 1:
                    break

                response.raise_for_status()

            data = response.json()
            try:
                return data["choices"][0]["message"]["content"].strip()
            except (KeyError, IndexError, TypeError):
                raise ValueError(f"Unexpected model response format: {json.dumps(data)[:500]}")

    if last_error is not None:
        last_error.raise_for_status()

    raise ValueError("Model call failed before receiving a response.")


def parse_json_output(text: str) -> dict:
    text = text.strip()
    if not text:
        raise ValueError("Model returned empty response when JSON was expected.")

    def _extract_json_candidate(raw: str) -> str:
        fenced = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", raw, re.IGNORECASE)
        if fenced:
            return fenced.group(1).strip()

        match = re.search(r"\{[\s\S]*\}", raw)
        if match:
            return match.group(0).strip()

        return raw.strip()

    def _sanitize_json_candidate(raw: str) -> str:
        # Escape invalid backslashes that often appear in LLM-generated JSON strings.
        sanitized = re.sub(r'(?<!\\)\\(?!["\\/bfnrtu])', r'\\\\', raw)
        # Remove trailing commas before closing braces/brackets.
        sanitized = re.sub(r",\s*([}\]])", r"\1", sanitized)
        return sanitized

    candidate = _extract_json_candidate(text)

    for attempt in (candidate, _sanitize_json_candidate(candidate)):
        try:
            return json.loads(attempt)
        except json.JSONDecodeError:
            continue

    raise ValueError("Could not parse valid JSON from model output.")


def normalize_features(raw_features):
    if isinstance(raw_features, list):
        return [str(item).strip() for item in raw_features if str(item).strip()]

    if isinstance(raw_features, str):
        tokens = re.split(r"[\n,]", raw_features)
        return [token.strip() for token in tokens if token.strip()]

    return []


def normalize_competitor_references(raw_refs, max_items: int = 8) -> list[dict]:
    if not isinstance(raw_refs, list):
        return []

    normalized = []
    for item in raw_refs[:max_items]:
        if not isinstance(item, dict):
            continue

        name = clamp_text(str(item.get("name", "")).strip(), 140)
        url = clamp_text(str(item.get("url", "")).strip(), 320)
        if not name or not url:
            continue

        line_one = clamp_text(
            str(item.get("problem") or item.get("relationToIdea") or "Relevant competitor/product in this market.").strip(),
            220,
        )
        line_two = clamp_text(
            str(item.get("solution") or item.get("relationToIdea") or "Included as a reference for feature and positioning analysis.").strip(),
            220,
        )

        normalized.append(
            {
                "name": name,
                "url": url,
                "lineOne": line_one,
                "lineTwo": line_two,
            }
        )

    return normalized


def append_competitor_references_section(srs_text: str, references: list[dict]) -> str:
    base = (srs_text or "").strip()
    if not references:
        return base

    lines = ["", "8. Competitor/Product References"]
    for idx, ref in enumerate(references, start=1):
        lines.append("")
        lines.append(f"{idx}. {ref.get('name', 'Unknown')}")
        lines.append(f"Link: {ref.get('url', '')}")
        lines.append(f"Summary: {ref.get('lineOne', '')}")
        lines.append(f"Details: {ref.get('lineTwo', '')}")

    appendix = "\n".join(lines).strip()
    if not base:
        return appendix

    return f"{base}\n\n{appendix}".strip()


def market_research_user_prompt(raw_idea: str, features: list[str]) -> str:
    raw_idea = clamp_text(raw_idea, 4500)
    features = compact_features(features, max_items=20, max_chars_per_item=180)
    feature_block = "\n".join([f"- {item}" for item in features]) if features else "- Not provided"
    return f"""INPUT:
Product Idea:
\"\"\"
{raw_idea}
\"\"\"

Current Features:
{feature_block}

OUTPUT REQUIREMENTS:
Provide structured information for 3-5 highly relevant competitors including:
- Product name
- Working URL
- Problem they solve
- Solution approach
- Main features (concise and meaningful)
- Relation to this idea
- Similarity score (0-10)

Also include:
- Market summary (concise landscape view)
- Differentiation factor (what makes this idea unique or competitive)
"""


def structure_analysis_user_prompt(raw_idea: str, features: list[str], research_text: str) -> str:
    raw_idea = clamp_text(raw_idea, 4500)
    features = compact_features(features, max_items=20, max_chars_per_item=180)
    research_text = clamp_text(research_text, 16000)
    return f"""INPUT:
Original Idea:
\"\"\"
{raw_idea}
\"\"\"

Features:
{json.dumps(features, ensure_ascii=True)}

Market Research:
\"\"\"
{research_text}
\"\"\"

OUTPUT REQUIREMENTS:
- Generate a clean and professional project title
- Extract a clear problem statement
- Extract a practical solution
- Structure competitor data consistently
- Include summary and differentiation factor

OUTPUT FORMAT:
Return ONLY valid JSON (no markdown, no explanation, no extra text):

{{
  \"title\": \"\",
  \"problem\": \"\",
  \"solution\": \"\",
  \"analysis\": {{
    \"competitors\": [
      {{
        \"name\": \"\",
        \"url\": \"\",
        \"similarityScore\": 0,
        \"problem\": \"\",
        \"solution\": \"\",
        \"mainFeatures\": [],
        \"relationToIdea\": \"\"
      }}
    ],
    \"summary\": \"\",
    \"differentiationFactor\": \"\"
  }}
}}
"""


def suggestions_user_prompt(title: str, problem: str, solution: str, summary: str, competitors: list[dict]) -> str:
    title = clamp_text(title, 300)
    problem = clamp_text(problem, 1800)
    solution = clamp_text(solution, 1800)
    summary = clamp_text(summary, 2400)
    competitors = compact_competitors(competitors, max_items=8, max_chars_per_field=420)
    return f"""INPUT:
Title: {title}
Problem: {problem}
Solution: {solution}
Market Summary: {summary}
Competitors:
{json.dumps(competitors, ensure_ascii=True, indent=2)}

OUTPUT REQUIREMENTS:
Generate exactly 5 suggestions.

OUTPUT:
Return ONLY valid JSON:

{{
  \"suggestions\": [
    {{
      \"id\": \"\",
      \"type\": \"feature | strategic | technical\",
      \"title\": \"\",
      \"description\": \"\",
      \"sourceInspiration\": \"\"
    }}
  ]
}}
"""


def srs_user_prompt(title: str, problem: str, solution: str, final_features: list[str]) -> str:
    title = clamp_text(title, 300)
    problem = clamp_text(problem, 2000)
    solution = clamp_text(solution, 2000)
    final_features = compact_features(final_features, max_items=30, max_chars_per_item=220)
    return f"""INPUT:
Title: {title}
Problem: {problem}
Solution: {solution}
Final Features:
{json.dumps(final_features, ensure_ascii=True, indent=2)}

REQUIRED TEMPLATE (must follow exactly):

MVP Requirements Document

1. Product Overview
Problem:
Solution:
MVP Goal:

2. Core Features

3. User Flow
Step-by-step product flow:

4. Functional Requirements
Inputs:
System behavior:
Outputs:

5. Basic Technical Architecture
Frontend:
Backend:
Database / APIs:

6. Non-Functional Requirements
Performance:
Security:

7. Phase Plan
Phase 1 (MVP):
Phase 2 (Future):

IMPORTANT:
- Output must be plain text only
- Do not use markdown syntax (#, ##, -, *, ```)
- Do not use code fences
- Fill every section with product-specific details
"""


def normalize_srs_output(text: str) -> str:
    cleaned = (text or "").strip()
    if not cleaned:
        return cleaned

    fenced = re.match(r"^```(?:markdown|md)?\s*([\s\S]*?)\s*```$", cleaned, re.IGNORECASE)
    if fenced:
        cleaned = fenced.group(1).strip()

    lines = []
    for line in cleaned.splitlines():
        line = line.rstrip()
        if not line:
            lines.append("")
            continue

        # Remove markdown heading markers if the model returns them.
        line = re.sub(r"^\s*#{1,6}\s*", "", line)
        # Convert markdown bullets to plain text labels.
        line = re.sub(r"^\s*[-*]\s+", "", line)
        lines.append(line)

    return "\n".join(lines).strip()


def markdown_to_docx(markdown_text: str, title: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading(title or "Software Requirements Specification", 0)

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif re.match(r"^[-*]\s+", stripped):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", stripped), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        else:
            doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _render_frontend(page_mode: str = "home", checkout_plan: str = "", next_path: str = ""):
    return render_template("index.html", page_mode=page_mode, checkout_plan=checkout_plan, next_path=next_path)


@app.get("/")
def index():
    return _render_frontend(page_mode="home")


@app.get("/auth")
def auth_page():
    next_path = _safe_next_path(request.args.get("next", ""), default="")
    return _render_frontend(page_mode="auth", next_path=next_path)


@app.get("/pricing")
def pricing_page():
    if not session.get("username"):
        return redirect(url_for("auth_page", next=request.path))
    return _render_frontend(page_mode="pricing")


@app.get("/checkout/<plan_id>")
def checkout_page(plan_id: str):
    if not session.get("username"):
        return redirect(url_for("auth_page", next=request.path))

    valid_plans = {"starter", "growth", "scale"}
    normalized = plan_id.strip().lower()
    if normalized not in valid_plans:
        normalized = "growth"
    return _render_frontend(page_mode="checkout", checkout_plan=normalized)


@app.get("/workspace")
def workspace_page():
    if not session.get("username"):
        return redirect(url_for("auth_page", next=request.path))
    return _render_frontend(page_mode="workspace")


@app.get("/api/health")
def health_check():
    return jsonify({"status": "ok"})


@app.get("/api/auth/session")
def auth_session():
    username = session.get("username")
    if not username:
        return jsonify({"authenticated": False, "user": None})

    users_data = _read_json(USERS_FILE, {"users": []})
    users = users_data.get("users", [])
    user = _find_user_by_username(users, username)
    if not user:
        session.clear()
        return jsonify({"authenticated": False, "user": None})

    return jsonify(
        {
            "authenticated": True,
            "user": {
                "username": user.get("username"),
                "isPremium": bool(user.get("isPremium", False)),
                "planId": user.get("planId", "free"),
                "ideaQuota": int(user.get("ideaQuota", 1)),
            },
        }
    )


@app.post("/api/auth/signup")
def auth_signup():
    try:
        payload = request.get_json(force=True)
        email = str(payload.get("email", "")).strip()
        username = str(payload.get("username", "")).strip()
        password = str(payload.get("password", "")).strip()

        if "@" not in email or len(email) < 5:
            return jsonify({"error": "A valid email is required."}), 400
        if len(username) < 3:
            return jsonify({"error": "Username must be at least 3 characters."}), 400
        if len(password) < 6:
            return jsonify({"error": "Password must be at least 6 characters."}), 400

        users_data = _read_json(USERS_FILE, {"users": []})
        users = users_data.get("users", [])

        if _find_user_by_username(users, username):
            return jsonify({"error": "This username already exists. Please choose another."}), 409

        if _find_user_by_email(users, email):
            return jsonify({"error": "This email is already registered."}), 409

        password_hash = _hash_password(password)
        if _is_duplicate_password(users, password_hash):
            return jsonify({"error": "This password is already used by another account. Use a different password."}), 409

        users.append(
            {
                "id": str(uuid.uuid4()),
                "email": email,
                "username": username,
                "passwordHash": password_hash,
                "isPremium": False,
                "planId": "free",
                "ideaQuota": 1,
                "createdAt": datetime.utcnow().isoformat(),
            }
        )
        users_data["users"] = users
        _write_json(USERS_FILE, users_data)

        session["username"] = username

        return jsonify(
            {
                "message": "Signup successful.",
                "user": {
                    "username": username,
                    "isPremium": False,
                    "planId": "free",
                    "ideaQuota": 1,
                },
            }
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.post("/api/auth/login")
def auth_login():
    try:
        payload = request.get_json(force=True)
        username = str(payload.get("username", "")).strip()
        password = str(payload.get("password", "")).strip()

        users_data = _read_json(USERS_FILE, {"users": []})
        users = users_data.get("users", [])
        user = _find_user_by_username(users, username)

        if not user:
            return jsonify({"error": "Invalid username or password."}), 401

        if user.get("passwordHash") != _hash_password(password):
            return jsonify({"error": "Invalid username or password."}), 401

        session["username"] = user.get("username")

        return jsonify(
            {
                "message": "Login successful.",
                "user": {
                    "username": user.get("username"),
                    "isPremium": bool(user.get("isPremium", False)),
                    "planId": user.get("planId", "free"),
                    "ideaQuota": int(user.get("ideaQuota", 1)),
                },
            }
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.post("/api/auth/logout")
def auth_logout():
    session.clear()
    return jsonify({"message": "Logged out."})


@app.get("/api/pricing")
def get_pricing():
    plans = [
        {
            "id": "starter",
            "name": "Starter",
            "pricePkr": 2500,
            "ideaQuota": 10,
            "description": "Best for solo founders validating focused concepts.",
            "features": [
                "10 complete idea validation runs",
                "Market analysis with competitor mapping",
                "Improvement recommendation workflow",
                "SRS generation and Word export",
                "Basic support",
            ],
        },
        {
            "id": "growth",
            "name": "Growth",
            "pricePkr": 5500,
            "ideaQuota": 35,
            "description": "For teams iterating multiple product directions.",
            "features": [
                "35 complete idea validation runs",
                "Advanced competitor intelligence",
                "Priority processing for generation steps",
                "SRS export for all approved ideas",
                "Team-friendly planning capacity",
                "Priority support",
            ],
        },
        {
            "id": "scale",
            "name": "Scale",
            "pricePkr": 9500,
            "ideaQuota": 100,
            "description": "For agencies and accelerators with high idea volume.",
            "features": [
                "100 complete idea validation runs",
                "High-volume multi-idea workflow",
                "Structured strategic improvement planning",
                "Full SRS export at scale",
                "Best value per idea",
                "Priority queue + support",
                "Designed for agency and accelerator operations",
            ],
        },
    ]
    return jsonify({"plans": plans})


@app.post("/api/payment/submit")
def submit_payment_proof():
    try:
        session_username = session.get("username")
        if not session_username:
            return jsonify({"error": "Please login first."}), 401

        username = str(request.form.get("username", "")).strip()
        trx_id = str(request.form.get("trxId", "")).strip()
        plan_id = str(request.form.get("planId", "")).strip().lower() or "growth"
        screenshot = request.files.get("screenshot")

        if not username:
            return jsonify({"error": "Username is required."}), 400
        if username.lower() != session_username.lower():
            return jsonify({"error": "Submitted username must match logged in account."}), 400
        if len(trx_id) < 4:
            return jsonify({"error": "Transaction ID is required."}), 400
        if screenshot is None or screenshot.filename == "":
            return jsonify({"error": "Screenshot or proof file is required."}), 400
        if not _allowed_payment_file(screenshot.filename):
            return jsonify({"error": "Allowed proof formats: png, jpg, jpeg, webp, pdf."}), 400

        plan_map = {
            "starter": {"ideaQuota": 10, "pricePkr": 2500},
            "growth": {"ideaQuota": 35, "pricePkr": 5500},
            "scale": {"ideaQuota": 100, "pricePkr": 9500},
        }
        selected_plan = plan_map.get(plan_id, plan_map["growth"])
        selected_plan_id = plan_id if plan_id in plan_map else "growth"

        safe_name = secure_filename(screenshot.filename)
        ext = os.path.splitext(safe_name)[1].lower()
        proof_filename = f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex}{ext}"
        proof_path = os.path.join(UPLOAD_DIR, proof_filename)
        screenshot.save(proof_path)

        users_data = _read_json(USERS_FILE, {"users": []})
        users = users_data.get("users", [])
        user = _find_user_by_username(users, username)
        if not user:
            return jsonify({"error": "User account not found."}), 404

        # Activate premium immediately on submission as requested.
        user["isPremium"] = True
        user["planId"] = selected_plan_id
        user["ideaQuota"] = int(selected_plan["ideaQuota"])
        user["premiumActivatedAt"] = datetime.utcnow().isoformat()
        _write_json(USERS_FILE, users_data)

        payments_data = _read_json(PAYMENTS_FILE, {"payments": []})
        payments = payments_data.get("payments", [])
        payments.append(
            {
                "id": str(uuid.uuid4()),
                "username": username,
                "trxId": trx_id,
                "planId": selected_plan_id,
                "pricePkr": selected_plan["pricePkr"],
                "proofFile": proof_filename,
                "proofPath": proof_path,
                "submittedAt": datetime.utcnow().isoformat(),
                "status": "received-premium-activated",
            }
        )
        payments_data["payments"] = payments
        _write_json(PAYMENTS_FILE, payments_data)

        return jsonify(
            {
                "message": "Payment proof received. Premium is now active.",
                "user": {
                    "username": user.get("username"),
                    "isPremium": True,
                    "planId": selected_plan_id,
                    "ideaQuota": int(selected_plan["ideaQuota"]),
                },
            }
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.post("/api/stage1")
def stage1_market_analysis():
    try:
        payload = request.get_json(force=True)
        raw_idea = str(payload.get("rawIdea", "")).strip()
        features = normalize_features(payload.get("features", []))

        if not raw_idea:
            return jsonify({"error": "rawIdea is required"}), 400

        config = get_client_config()

        research_text = call_chat_model(
            model=config["model"],
            system_prompt=MARKET_RESEARCH_SYSTEM_PROMPT,
            user_prompt=market_research_user_prompt(raw_idea, features),
            temperature=0.2,
        )

        structured_raw = call_chat_model(
            model=config["model"],
            system_prompt=STRUCTURE_ANALYSIS_SYSTEM_PROMPT,
            user_prompt=structure_analysis_user_prompt(raw_idea, features, research_text),
            temperature=0.1,
        )

        structured = parse_json_output(structured_raw)

        return jsonify(
            {
                "researchText": research_text,
                "structured": structured,
            }
        )

    except requests.HTTPError as exc:
        details = ""
        try:
            details = exc.response.text[:1000] if exc.response is not None else ""
        except Exception:
            pass
        return jsonify({"error": "LLM API request failed", "details": details}), 502
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.post("/api/stage2")
def stage2_improvements():
    try:
        payload = request.get_json(force=True)

        title = str(payload.get("title", "")).strip()
        problem = str(payload.get("problem", "")).strip()
        solution = str(payload.get("solution", "")).strip()
        summary = str(payload.get("summary", "")).strip()
        competitors = payload.get("competitors", [])

        if not title:
            return jsonify({"error": "title is required"}), 400

        config = get_client_config()

        suggestions_raw = call_chat_model(
            model=config["model"],
            system_prompt=SUGGESTIONS_SYSTEM_PROMPT,
            user_prompt=suggestions_user_prompt(title, problem, solution, summary, competitors),
            temperature=0.4,
        )

        suggestions = parse_json_output(suggestions_raw)
        return jsonify(suggestions)

    except requests.HTTPError as exc:
        details = ""
        try:
            details = exc.response.text[:1000] if exc.response is not None else ""
        except Exception:
            pass
        return jsonify({"error": "LLM API request failed", "details": details}), 502
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.post("/api/stage3")
def stage3_generate_srs():
    try:
        payload = request.get_json(force=True)

        title = str(payload.get("title", "")).strip()
        problem = str(payload.get("problem", "")).strip()
        solution = str(payload.get("solution", "")).strip()
        final_features = normalize_features(payload.get("finalFeatures", []))
        competitor_references = normalize_competitor_references(payload.get("competitorReferences", []))

        if not title:
            return jsonify({"error": "title is required"}), 400

        config = get_client_config()

        srs_markdown = call_chat_model(
            model=config["model"],
            system_prompt=SRS_SYSTEM_PROMPT,
            user_prompt=srs_user_prompt(title, problem, solution, final_features),
            temperature=0.2,
        )

        srs_markdown = normalize_srs_output(srs_markdown)
        srs_markdown = append_competitor_references_section(srs_markdown, competitor_references)

        return jsonify({"srsMarkdown": srs_markdown})

    except requests.HTTPError as exc:
        details = ""
        try:
            details = exc.response.text[:1000] if exc.response is not None else ""
        except Exception:
            pass
        return jsonify({"error": "LLM API request failed", "details": details}), 502
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.post("/api/download")
def download_srs_docx():
    try:
        payload = request.get_json(force=True)
        srs_markdown = str(payload.get("srsMarkdown", "")).strip()
        title = str(payload.get("title", "Software Requirements Specification")).strip()

        if not srs_markdown:
            return jsonify({"error": "srsMarkdown is required"}), 400

        file_buffer = markdown_to_docx(srs_markdown, title)

        safe_title = re.sub(r"[^a-zA-Z0-9_-]+", "_", title).strip("_") or "srs"
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        file_name = f"{safe_title}_{timestamp}.docx"

        return send_file(
            file_buffer,
            as_attachment=True,
            download_name=file_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
